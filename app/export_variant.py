"""
Экспорт варианта ЕГЭ/ОГЭ в PDF и DOCX.

Что умеет:
  • собирает случайный вариант из таблицы tasks;
  • экспортирует вариант в PDF в стиле реального КИМ;
  • экспортирует вариант в DOCX;
  • умеет создавать версию без ответов и версию с ответами;
  • вставляет локальные изображения из image_path, если файл найден;
  • корректно работает с темами, question, answer, exam_type, task_number.

Улучшения по сравнению с предыдущей версией:
  • Шапка с реквизитами: ФИО, класс, школа, дата, номер варианта
  • Инструкции по ФИПИ: время, что разрешено, система баллов
  • Баллы рядом с номером каждого задания
  • Клеточки для ответа вместо строки «Ответ: _____»
  • Для ОГЭ Часть 2 — разрыв страницы + инструкция «развёрнутый ответ»
  • Таблица баллов в конце (при --answers)
  • Бланк ответов Части 1 как отдельная секция

Запуск из папки app/:
  python export_variant.py --exam ege_profile --pdf
  python export_variant.py --exam ege_base --pdf --answers
  python export_variant.py --exam oge --pdf --docx --answers

Примеры:
  python export_variant.py --exam ege_profile --pdf --variant-number 331
  python export_variant.py --exam ege_base --docx --outdir exports

Зависимости:
  pip install reportlab python-docx pillow
"""

from __future__ import annotations

import argparse
import os
import random
import re
import sqlite3
from dataclasses import dataclass
from xml.sax.saxutils import escape as xml_escape
from datetime import datetime
from pathlib import Path
from typing import Iterable

# PDF
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT, TA_RIGHT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm, mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    Image,
    KeepTogether,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
    HRFlowable,
)

# DOCX
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ════════════════════════════════════════════════════════════════════════════
# Пути
# ════════════════════════════════════════════════════════════════════════════

APP_DIR = Path(__file__).resolve().parent
BASE_DIR = APP_DIR.parent
DB_PATH = APP_DIR / "bot_database.db"
DEFAULT_OUTDIR = APP_DIR / "exports"
REFERENCE_DIR = APP_DIR / "reference_materials"

REFERENCE_PDF_CANDIDATES = {
    "oge": [
        REFERENCE_DIR / "oge_reference.pdf",
        APP_DIR / "Справочные материалы ОГЭ 2022.pdf",
        BASE_DIR / "Справочные материалы ОГЭ 2022.pdf",
    ],
    "ege_base": [
        REFERENCE_DIR / "ege_base_reference.pdf",
        APP_DIR / "Справочный материал (БАЗА) (1).pdf",
        BASE_DIR / "Справочный материал (БАЗА) (1).pdf",
    ],
}


# ════════════════════════════════════════════════════════════════════════════
# Настройки вариантов
# ════════════════════════════════════════════════════════════════════════════

EXAM_TITLES = {
    "oge":        "ОГЭ по математике",
    "ege_base":   "ЕГЭ по математике · базовый уровень",
    "ege_profile":"ЕГЭ по математике · профильный уровень",
}

EXAM_SHORT_TITLES = {
    "oge":        "ОГЭ",
    "ege_base":   "ЕГЭ база",
    "ege_profile":"ЕГЭ профиль",
}

EXAM_RANGES = {
    "oge":        range(1, 26),
    "ege_base":   range(1, 22),
    "ege_profile":range(1, 20),
}

# Время выполнения (минуты)
EXAM_TIME = {
    "oge":        235,
    "ege_base":   180,
    "ege_profile":235,
}

PROFILE_PART1_MAX = 12
PROFILE_PART2_START = 13

# Баллы за задания
TASK_SCORES = {
    "oge": {
        **{n: 1 for n in range(1, 20)},   # задания 1–19: по 1 баллу
        **{n: 2 for n in range(20, 26)},   # задания 20–25: по 2 балла
    },
    "ege_base": {n: 1 for n in range(1, 22)},
    "ege_profile": {
        **{n: 1 for n in range(1, 13)},    # часть 1: по 1 баллу
        **{n: 2 for n in range(13, 20)},   # часть 2: по 2 балла
    },
}

TOPIC_NAMES = {
    # ЕГЭ профиль
    "ege_p01_planimetry":           "Планиметрия",
    "ege_p02_vectors":              "Векторы",
    "ege_p03_stereometry":          "Стереометрия",
    "ege_p04_probability":          "Теория вероятностей",
    "ege_p05_probability_complex":  "Теория вероятностей: сложные задачи",
    "ege_p06_stereometry":          "Стереометрия",
    "ege_p07_expr":                 "Значение выражения",
    "ege_p08_equations":            "Уравнения",
    "ege_p09_word_problems":        "Текстовые задачи",
    "ege_p10_text_problems":        "Текстовые задачи",
    "ege_p11_functions":            "Функции и графики",
    "ege_p12_derivative":           "Производная и исследование функций",
    "ege_p13_equations":            "Уравнения",
    "ege_p14_stereometry":          "Стереометрия",
    "ege_p15_inequalities":         "Неравенства",
    "ege_p16_planimetry":           "Планиметрия",
    "ege_p17_economics":            "Экономическая задача",
    "ege_p18_parameters":           "Задача с параметром",
    "ege_p19_numbers":              "Числа и их свойства",
    # ЕГЭ база
    "ege_b01_units":                "Единицы измерения",
    "ege_b02_reading_graphs":       "Чтение графиков и диаграмм",
    "ege_b03_tables_graphs":        "Таблицы, графики и диаграммы",
    "ege_b04_probability":          "Теория вероятностей",
    "ege_b05_practical_calculus":   "Прикладные задачи",
    "ege_b06_geometry":             "Планиметрия",
    "ege_b07_graphs":               "Графики и диаграммы",
    "ege_b08_statements":           "Выбор верных утверждений",
    "ege_b09_calculations":         "Вычисления",
    "ege_b10_planimetry":           "Прикладная планиметрия",
    "ege_b11_stereometry":          "Стереометрия",
    "ege_b12_functions":            "Исследование функций",
    "ege_b13_equations":            "Уравнения",
    "ege_b14_inequalities":         "Неравенства",
    "ege_b15_finance":              "Финансовая математика",
    "ege_b16_progressions":         "Прогрессии",
    "ege_b17_geometry":             "Геометрия",
    "ege_b18_ineq":                 "Числа и неравенства",
    "ege_b18_inequalities":         "Числа и неравенства",
    "ege_b19_digits":               "Числа и цифры",
    "ege_b20_text_problems":        "Текстовые задачи",
    "ege_b21_logic":                "Задачи на смекалку",
    # ОГЭ
    "general":                      "Общее",
    "probability":                  "Теория вероятностей",
    "degrees":                      "Степени",
    "arithmetic_square_root":       "Арифметические корни",
    "ordinary_fractions":           "Обыкновенные дроби",
    "decimal_fractions":            "Десятичные дроби",
    "triangles":                    "Треугольники",
    "circles":                      "Окружности",
    "parallelogram":                "Параллелограмм",
    "trapezoid":                    "Трапеция",
    "rectangle":                    "Прямоугольник",
    "rhombus":                      "Ромб",
    "square":                       "Квадрат",
    "grid":                         "Клетчатая плоскость",
    "arithmetic_progression":       "Арифметическая прогрессия",
    "geometric_progression":        "Геометрическая прогрессия",
    "linear_function":              "Линейная функция",
    "quadratic_function":           "Квадратичная функция",
    "hyperbola":                    "Обратная пропорциональность",
    "mixed_graphs":                 "Смешанные графики",
    "n20_systems":                  "Системы уравнений",
    "n20_equations":                "Уравнения",
    "n20_inequalities":             "Неравенства",
    "inequalities":                 "Неравенства",
    "n21_motion_line":              "Движение по прямой",
    "n21_motion_water":             "Движение по воде",
    "n21_work":                     "Работа",
    "n21_percent":                  "Проценты, смеси и сплавы",
    "n23_circle":                   "Окружность",
    "n23_parallelogram":            "Параллелограмм",
    "n23_rhombus":                  "Ромб",
    "n23_triangle":                 "Треугольник",
}


@dataclass
class Task:
    id: int
    exam_type: str
    task_number: int
    topic: str
    question: str | None
    answer: str
    image_path: str | None = None
    set_id: int | None = None

    @property
    def topic_title(self) -> str:
        return TOPIC_NAMES.get(self.topic, self.topic.replace("_", " ").strip())

    @property
    def score(self) -> int:
        return TASK_SCORES.get(self.exam_type, {}).get(self.task_number, 1)

    @property
    def is_part2(self) -> bool:
        """Задание второй части (развёрнутый ответ)."""
        if self.exam_type == "oge":
            return self.task_number >= 20
        if self.exam_type == "ege_profile":
            return self.task_number >= PROFILE_PART2_START
        return False


# ════════════════════════════════════════════════════════════════════════════
# Работа с БД
# ════════════════════════════════════════════════════════════════════════════

def connect_db(db_path: Path = DB_PATH) -> sqlite3.Connection:
    if not db_path.exists():
        raise FileNotFoundError(f"База данных не найдена: {db_path}")
    conn = sqlite3.connect(db_path)
    conn.row_factory = sqlite3.Row
    return conn


def row_to_task(row: sqlite3.Row) -> Task:
    return Task(
        id=row["id"],
        exam_type=row["exam_type"],
        task_number=row["task_number"],
        topic=row["topic"],
        question=row["question"],
        answer=str(row["answer"]),
        image_path=row["image_path"],
        set_id=row["set_id"] if "set_id" in row.keys() else None,
    )


def fetch_random_task(conn: sqlite3.Connection, exam_type: str, task_number: int) -> Task | None:
    row = conn.execute(
        """
        SELECT * FROM tasks
        WHERE exam_type = ? AND task_number = ?
        ORDER BY RANDOM()
        LIMIT 1
        """,
        (exam_type, task_number),
    ).fetchone()
    return row_to_task(row) if row else None


def fetch_oge_set_1_5(conn: sqlite3.Connection) -> list[Task]:
    rows = conn.execute(
        """
        SELECT DISTINCT topic, set_id FROM tasks
        WHERE exam_type='oge'
          AND task_number BETWEEN 1 AND 5
          AND set_id IS NOT NULL
        """
    ).fetchall()

    pairs = [(r["topic"], r["set_id"]) for r in rows]
    if pairs:
        topic, set_id = random.choice(pairs)
        set_rows = conn.execute(
            """
            SELECT * FROM tasks
            WHERE exam_type='oge'
              AND task_number BETWEEN 1 AND 5
              AND topic=? AND set_id=?
            ORDER BY task_number
            """,
            (topic, set_id),
        ).fetchall()
        return [row_to_task(r) for r in set_rows]

    tasks: list[Task] = []
    for n in range(1, 6):
        t = fetch_random_task(conn, "oge", n)
        if t:
            tasks.append(t)
    return tasks


def build_random_variant(exam_type: str, db_path: Path = DB_PATH) -> list[Task]:
    if exam_type not in EXAM_RANGES:
        raise ValueError(f"Неизвестный exam_type: {exam_type}")

    conn = connect_db(db_path)
    try:
        tasks: list[Task] = []

        if exam_type == "oge":
            tasks.extend(fetch_oge_set_1_5(conn))
            start = 6
        else:
            start = 1

        for task_number in EXAM_RANGES[exam_type]:
            if task_number < start:
                continue
            task = fetch_random_task(conn, exam_type, task_number)
            if task:
                tasks.append(task)

        tasks.sort(key=lambda t: t.task_number)
        return tasks
    finally:
        conn.close()


# ════════════════════════════════════════════════════════════════════════════
# Очистка текста
# ════════════════════════════════════════════════════════════════════════════

PDF_GLYPH_REPLACEMENTS = {
    "\uf028": "(", "\uf029": ")", "\uf02d": "-", "\uf02b": "+",
    "\uf03d": "=", "\uf03c": "<", "\uf03e": ">",
    "\uf0b0": "°", "\uf0d7": "×", "\uf0f7": ":",
    "\uf061": "α", "\uf062": "β", "\uf067": "γ",
    "\uf070": "π", "\uf0d0": "∠", "\uf0a3": "≤", "\uf0b3": "≥",
    "\uf0ce": "∞", "\uf0c5": "∈", "\uf0c7": "∪",
    "\uf0bc": "1/4", "\uf0bd": "1/2", "\uf0be": "3/4",
    "\uf0a2": "√", "\uf0d6": "√",
    "\uf0a0": " ",
    "?": "",
}

BOX_CHARS = "□▯◻◼■▫▪⬚"


def restore_box_variables(text: str) -> str:
    """
    Исправляет квадратики/служебные символы после импорта из PDF.

    Для обычных алгебраических заданий квадратик часто означает x.
    Для тригонометрии ЕГЭ профиль №7 квадратик часто появляется вместо α:
        cos□ = ...  -> cosα = ...
        □ ∈ (...)   -> α ∈ (...)
    Для задания 19 ОГЭ □□□ + □□ + □ — содержательные клеточки, не трогаем.
    """
    box = re.escape(BOX_CHARS)

    # Задание 19 ОГЭ: □□□ + □□ + □ — клеточки для цифр, смысловой элемент.
    if re.search(rf"[{box}]{{2,}}\s*\+\s*[{box}]", text):
        return text

    # Тригонометрия: квадрат вместо α.
    text = re.sub(rf"\b(sin|cos|tg|tan|ctg)\s*[{box}]+", r"\1α", text, flags=re.IGNORECASE)
    text = re.sub(rf"[{box}]+\s*(?=∈)", "α", text)
    text = re.sub(r"\b(sin|cos|tg|tan|ctg)\s*(?==)", r"\1α", text, flags=re.IGNORECASE)

    # Алгебра: квадрат вместо x.
    text = re.sub(rf"(?<=\d)[{box}]+", "x", text)
    text = re.sub(rf"(?:(?<=^)|(?<=[\s(=+\-*/:<>≤≥]))[{box}]+(?=\s*(?:[+\-*/:=<>≤≥)]|$))", "x", text)

    # Оставшиеся служебные квадратики убираем, чтобы они не печатались в PDF.
    text = re.sub(rf"[{box}]+", "", text)
    return text

def clean_text(text: str | None) -> str:
    if not text:
        return ""
    text = str(text)
    for bad, good in PDF_GLYPH_REPLACEMENTS.items():
        text = text.replace(bad, good)
    text = re.sub(r"[∠∡]\s*([A-ZА-ЯЁ][A-ZА-ЯЁ0-9₁₂₃₄₅₆₇₈₉]*)", r"угол \1", text)
    text = text.replace("∠", "угол ").replace("∡", "угол ")
    # После строки с заменой ∠ на "угол":
    text = re.sub(r'\bsin\s*угол\b', 'sin угла', text, flags=re.IGNORECASE)
    text = re.sub(r'\bcos\s*угол\b', 'cos угла', text, flags=re.IGNORECASE)
    text = restore_box_variables(text)
    # Дополнительная защита для тригонометрии ЕГЭ профиль: после OCR иногда теряется α.
    text = re.sub(r"\b(sin|cos|tg|tan|ctg)\s*=", r"\1α =", text, flags=re.IGNORECASE)
    text = re.sub(r"(?<![A-Za-zА-Яа-яαβγ])\s+∈", " α ∈", text)
    text = text.replace("\u00a0", " ").replace("\u202f", " ")
    text = text.replace("−", "-").replace("–", "-").replace("—", "-")
    text = text.replace("⋅", "·").replace("∙", "·")
    text = text.replace("^2", "²").replace("^3", "³")
    text = re.sub(r"(?:(?<=уравнение )|(?<=[(=\s+\-*/:]))²", "x²", text)
    text = re.sub(r"(?:(?<=уравнение )|(?<=[(=\s+\-*/:]))³", "x³", text)
    if re.search(r"(?:^|[\s(=+\-*/:])x?[²³]", text) and re.search(r"\s+x\s*$", text):
        text = re.sub(r"\s+x\s*$", "", text)
    text = re.sub(r"(?<![;,])\s+([1-9]\))", r"\n\1", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\s+([,.;:!?])", r"\1", text)
    text = re.sub(r"\n\s+", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)

    return text.strip()


def pdf_text(text: str | None) -> str:
    return xml_escape(clean_text(text)).replace("\n", "<br/>")


def pdf_no_wrap_math_fragment(pdf_s: str, task_number: int | None = None) -> str:
    """
    Для ЕГЭ профиль №6 и №12 не даёт математическому выражению
    разрываться посередине строки.

    №6: обычно выражение стоит после двоеточия:
        Найдите корень уравнения: (3x - 7)² = (3x + 1)²

    №12: обычно выражение стоит после слова «функции» и до «на отрезке»:
        Найдите наименьшее значение функции y = 12x - 12ln(x + 2) + 7 на отрезке ...
    """
    def wrap_expr(expr: str) -> str:
        expr = expr.strip()
        if not expr:
            return expr
        # Неразрывные пробелы внутри самой формулы.
        expr = expr.replace(" ", "&#160;")
        return f"<nobr>{expr}</nobr>"

    # №6 и похожие: формула после двоеточия.
    if ":" in pdf_s:
        before, after = pdf_s.split(":", 1)
        after = after.strip()
        if after:
            # Оборачиваем только уравнение/выражение — до первой точки или скобки пояснения.
            # Остаток (пояснение вроде «Если уравнение имеет...») оставляем как есть.
            period_match = re.search(r"\.\s+[А-ЯЁA-Z]", after)
            if period_match:
                expr = after[:period_match.start() + 1]
                tail = after[period_match.start() + 1:]
            else:
                expr = after
                tail = ""
            return before + ": " + wrap_expr(expr) + tail

    # ОГЭ №8 и похожие: формула после слова «выражения».
    m = re.search(r"(?i)(выражени[яе]\s+)(.*)", pdf_s)
    if m:
        before = pdf_s[:m.start(2)]
        expr = m.group(2)
        if expr:
            return before + wrap_expr(expr)

    # №12: формула между "функции" и "на отрезке".
    if task_number == 12:
        m = re.search(r"(?i)(функции\s+)(.*?)(\s+на\s+отрезке\b.*)", pdf_s)
        if m:
            before = pdf_s[:m.start(2)]
            expr = m.group(2)
            after = pdf_s[m.end(2):]
            return before + wrap_expr(expr) + after

    return pdf_s



def answer_for_display(answer: str | None) -> str:
    ans = clean_text(answer)
    if "|" in ans:
        return " или ".join(part.strip() for part in ans.split("|") if part.strip())
    return ans


# ════════════════════════════════════════════════════════════════════════════
# Справочные материалы
# ════════════════════════════════════════════════════════════════════════════

def find_reference_pdf(exam_type: str) -> Path | None:
    for candidate in REFERENCE_PDF_CANDIDATES.get(exam_type, []):
        if candidate.exists() and candidate.is_file():
            return candidate
    return None


def prepend_reference_pdf(main_pdf: Path, reference_pdf: Path) -> Path:
    try:
        from pypdf import PdfReader, PdfWriter
    except Exception:
        print("⚠️ pypdf не установлен: справочные материалы не добавлены. pip install pypdf")
        return main_pdf
    try:
        writer = PdfWriter()
        for pdf in (reference_pdf, main_pdf):
            reader = PdfReader(str(pdf))
            for page in reader.pages:
                writer.add_page(page)
        tmp = main_pdf.with_suffix(".tmp.pdf")
        with tmp.open("wb") as f:
            writer.write(f)
        tmp.replace(main_pdf)
    except Exception as e:
        print(f"⚠️ Не удалось добавить справочные материалы: {e}")
    return main_pdf


# ════════════════════════════════════════════════════════════════════════════
# Изображения
# ════════════════════════════════════════════════════════════════════════════

def resolve_image_path(image_path: str | None) -> Path | None:
    if not image_path:
        return None
    if "::" in image_path:
        image_path = image_path.split("::")[0]
    if not image_path:
        return None
    if image_path.startswith(("photo", "doc", "http://", "https://")):
        return None
    p = Path(image_path)
    candidates = []
    if p.is_absolute():
        candidates.append(p)
    else:
        candidates.append(BASE_DIR / image_path)
        candidates.append(APP_DIR / image_path)
        if image_path.startswith("app/") or image_path.startswith("app\\"):
            candidates.append(BASE_DIR / image_path)
        else:
            candidates.append(APP_DIR / image_path)
            candidates.append(BASE_DIR / "app" / image_path)
    for candidate in candidates:
        if candidate.exists() and candidate.is_file():
            return candidate
    return None


def resolve_image_paths(image_path: str | None) -> list[Path]:
    if not image_path:
        return []
    paths: list[Path] = []
    for part in str(image_path).split("|"):
        p = resolve_image_path(part.strip())
        if p:
            paths.append(p)
    return paths


def get_export_image_size(task: Task, target: str = "pdf"):
    # ЕГЭ профиль №9: картинки/схемы в PDF делаем компактнее.
    if task.exam_type == "ege_profile" and task.task_number == 9:
        return (5.8 * cm, 3.8 * cm) if target == "pdf" else Cm(5.5)
    if task.exam_type == "oge":
        if task.task_number == 22:
            return (4.8 * cm, 2.7 * cm) if target == "pdf" else Cm(4.8)
        if task.task_number == 20:
            return (3.8 * cm, 2.1 * cm) if target == "pdf" else Cm(3.8)
        if task.task_number == 1:
            return (9.5 * cm, 6.5 * cm) if target == "pdf" else Cm(8.5)
        if task.task_number == 5:
            return (9.0 * cm, 6.2 * cm) if target == "pdf" else Cm(8.0)
        if task.task_number == 8:
            return (7.8 * cm, 5.5 * cm) if target == "pdf" else Cm(6.8)
        if task.task_number == 11:
            return (8.0 * cm, 6.0 * cm) if target == "pdf" else Cm(6.0)
        if task.task_number == 12:
            return (3.0 * cm, 2.0 * cm) if target == "pdf" else Cm(6.0)
        if task.task_number == 13:
            return (4.0 * cm, 2.8 * cm) if target == "pdf" else Cm(4.0)
            return (3.0 * cm, 2.0 * cm) if target == "pdf" else Cm(6.0)
        if task.task_number in (15, 16, 17, 18):
            return (6.5 * cm, 3.6 * cm) if target == "pdf" else Cm(6.5)
    return (7.8 * cm, 4.0 * cm) if target == "pdf" else Cm(7.8)


# ════════════════════════════════════════════════════════════════════════════
# Вспомогательные функции
# ════════════════════════════════════════════════════════════════════════════

def make_output_name(exam_type: str, variant_number: int | None, answers: bool, ext: str) -> str:
    print(f"[DEBUG] make_output_name: variant_number={variant_number!r}")  # ← добавить
    stamp = datetime.now().strftime("%Y%m%d_%H%M")
    exam_names = {"oge": "ОГЭ", "ege_base": "ЕГЭ_база", "ege_profile": "ЕГЭ_профиль"}
    exam = exam_names.get(exam_type, exam_type.replace("_", "-"))
    number = str(variant_number) if variant_number is not None else datetime.now().strftime("%H%M%S")
    suffix = "_с_ответами" if answers else ""
    return f"{exam}_вариант_{number}{suffix}.{ext}"

def get_exam_instruction(exam_type: str) -> str:
    """Инструкция по ФИПИ с временем и правилами."""
    minutes = EXAM_TIME.get(exam_type, 180)
    hours = minutes // 60
    mins = minutes % 60
    time_str = f"{hours} часа {mins} минут" if mins else f"{hours} часа"

    if exam_type == "oge":
        return (
            f"На выполнение экзаменационной работы отводится {time_str} ({minutes} минут). "
            "Работа состоит из двух частей: часть 1 — 19 заданий с кратким ответом (задания 1–19), "
            "часть 2 — 6 заданий с развёрнутым ответом (задания 20–25). "
            "Ответом к заданиям части 1 является целое число или конечная десятичная дробь. "
            "При выполнении заданий разрешается пользоваться справочными материалами. "
            "Использование калькулятора и мобильного телефона не допускается. "
            "Максимальный балл — 32 (19 баллов за часть 1 + 13 баллов за часть 2)."
        )
    if exam_type == "ege_base":
        return (
            f"На выполнение экзаменационной работы отводится {time_str} ({minutes} минут). "
            "Работа состоит из 21 задания с кратким ответом. "
            "Ответом является целое число, конечная десятичная дробь или последовательность цифр. "
            "Единицы измерения писать не нужно. "
            "При выполнении заданий разрешается пользоваться справочными материалами. "
            "Использование калькулятора и мобильного телефона не допускается. "
            "Максимальный балл — 21."
        )
    # ege_profile
    return (
        f"На выполнение экзаменационной работы отводится {time_str} ({minutes} минут). "
        "Часть 1 содержит 12 заданий с кратким ответом (целое число или конечная десятичная дробь). "
        "Часть 2 содержит 7 заданий с развёрнутым ответом — требуется полное обоснование. "
        "При выполнении заданий разрешается пользоваться справочными материалами. "
        "Использование калькулятора не допускается. "
        "Максимальный балл — 32."
    )


def get_part_title(exam_type: str, task_number: int) -> str | None:
    if exam_type == "ege_profile":
        return "Часть 1" if task_number <= PROFILE_PART1_MAX else "Часть 2"
    if exam_type == "oge":
        return "Часть 1" if task_number <= 19 else "Часть 2"
    return None


def repair_question_for_export(task: Task) -> str:
    q = clean_text(task.question)
    if task.exam_type != "oge":
        return q
    if task.task_number == 9:
        q = re.sub(r"(уравнени[ея]\s+)²", r"\1x²", q, flags=re.IGNORECASE)
        if "x" not in q.lower():
            q = re.sub(
                r"(уравнени[ея]\s+)([-+]?\d+(?:[.,]\d+)?)\s*([+\-])\s*(\d+(?:[.,]\d+)?)\s*=\s*([-+]?\d+(?:[.,]\d+)?)(?=\.|$)",
                lambda m: f"{m.group(1)}{m.group(2)}x {m.group(3)} {m.group(4)} = {m.group(5)}x",
                q, flags=re.IGNORECASE,
            )
    if task.task_number == 13:
        q = re.sub(r"(неравенств[ао]\s+)²", r"\1x²", q, flags=re.IGNORECASE)
        parts = re.split(r"(\n?1\))", q, maxsplit=1)
        condition = parts[0]
        tail = "".join(parts[1:]) if len(parts) > 1 else ""
        if "x" not in condition.lower():
            condition = re.sub(
                r"([-+]?\d+(?:[.,]\d+)?)\s*([+\-])\s*(\d+(?:[.,]\d+)?)\s*([<>≤≥]=?)\s*([-+]?\d+(?:[.,]\d+)?)\s*([+\-])\s*(\d+(?:[.,]\d+)?)",
                lambda m: f"{m.group(1)}x {m.group(2)} {m.group(3)} {m.group(4)} {m.group(5)}x {m.group(6)} {m.group(7)}",
                condition,
            )
            condition = re.sub(
                r"([-+]?\d+(?:[.,]\d+)?)\s*([+\-])\s*(\d+(?:[.,]\d+)?)\s*([<>≤≥]=?)\s*([-+]?\d+(?:[.,]\d+)?)(?=\.|$)",
                lambda m: f"{m.group(1)}x {m.group(2)} {m.group(3)} {m.group(4)} {m.group(5)}",
                condition,
            )
        q = condition + tail
    return q


def clean_question_for_variant(task: Task) -> str:
    q = repair_question_for_export(task) or "Решите задание по рисунку."

    # Убираем служебный префикс "Задание N ЕГЭ (...). " из текста задания
    q = re.sub(
        r"^Задание\s+\d+\s+ЕГЭ\s*\([^)]*\)\.\s*",
        "",
        q,
        flags=re.IGNORECASE,
    )

    # ЕГЭ профиль №16: условия через "; •" форматируем как список с новой строки
    if task.exam_type == "ege_profile" and task.task_number == 16:
        q = re.sub(r"\s*;\s*•\s*", "\n• ", q)
        q = re.sub(r"^•\s*", "• ", q)  # первый буллет если есть

    # ОГЭ №8: убираем подсказки вида «\nПодсказка: ...»
    if task.exam_type == "oge" and task.task_number == 8:
        q = re.sub(r"\n?Подсказка\s*:.*$", "", q,
                   flags=re.IGNORECASE | re.DOTALL).strip()

    # ОГЭ №20: убираем подсказки и примеры ввода.
    if task.exam_type == "oge" and task.task_number == 20:
        for marker in [r"\s*Введи\b", r"\s*Ввод\s+ответа\s*[:\-—–]", r"\s*Пример\s*:"]:
            q = re.split(marker, q, maxsplit=1, flags=re.IGNORECASE)[0]
        q = re.sub(
            r"\s*\(\s*Подсказка\s*:.*?(?=\s*Ввод\s+ответа|\s*$)",
            " ",
            q,
            flags=re.IGNORECASE | re.DOTALL,
        )


    # ЕГЭ база №7: в PDF/DOCX не выводим OCR-таблицы из текста.
    # Оставляем только условие + картинку, потому что соответствия/графики есть на изображении.
    if task.exam_type == "ege_base" and task.task_number == 7 and task.image_path:
        q = re.sub(
            r"\s*Е\.\s*А\.\s*Ширяева.*?\(тренаж[её]р\)\s*",
            " ",
            q,
            flags=re.IGNORECASE | re.DOTALL,
        )

        # Исправляем частые OCR-разрывы формул.
        q = re.sub(r"=\s*\n?\s*\(\s*\)\s*\n?\s*y\s*\n?\s*f\s*x", "= y = f(x)", q, flags=re.IGNORECASE)
        q = re.sub(r"f\s*\(\s*\)\s*y\s*f\s*x\s*\.?", "y = f(x).", q, flags=re.IGNORECASE)
        q = re.sub(r"вида\s*\n?\s*=\s*\n?\s*y\s*\n?\s*kx\s*\+\s*b\s*\.?", "вида y = kx + b.", q, flags=re.IGNORECASE)
        q = re.sub(r"вида\s*\n?\s*2\s*\n?\s*=\s*\n?\s*y\s*\n?\s*ax\s*\+\s*bx\s*\+\s*c\s*\.?", "вида y = ax² + bx + c.", q, flags=re.IGNORECASE)

        # Всё, что начинается с текстовой расшифровки таблицы, убираем.
        # Важно: режем ТОЛЬКО по началу строки (\n + ключевое слово),
        # иначе «изображены графики функций» обрезается посередине фразы.
        q = re.split(
            r"\n\s*(?:ИНТЕРВАЛЫ\s+ВРЕМЕНИ|ИНТЕРВАЛЫ|ПЕРИОДЫ\s+ВРЕМЕНИ|ПЕРИОДЫ|"
            r"ГРАФИКИ\s+ФУНКЦИЙ|ГРАФИКИ\s+ПРОИЗВОДНЫХ|ГРАФИКИ|ТОЧКИ|"
            r"ФУНКЦИИ|МЕСЯЦЫ|ХАРАКТЕРИСТИКИ|КОЭФФИЦИЕНТЫ|"
            r"УГЛОВЫЕ\s+КОЭФФИЦИЕНТЫ|ЗНАЧЕНИЯ\s+ПРОИЗВОДНОЙ)\b",
            q,
            maxsplit=1,
            flags=re.IGNORECASE,
        )[0].strip()

        # Схлопываем ВСЕ пробельные символы включая одиночные \n от OCR
        # (иначе "цена\nакции\nв\nрублях" даёт <br/> в PDF после каждого слова).
        q = re.sub(r"\s+", " ", q).strip()
        if q and not q.endswith((".", "?", "!")):
            q += "."


    # ЕГЭ база №18: если есть изображение, пункты выводим с новой строки.
    if task.exam_type == "ege_base" and task.task_number == 18 and task.image_path:
        q = re.sub(r"\s+([АБВГ]\))", r"\n\1", q)
        q = re.sub(r"\s+([1-4]\))", r"\n\1", q)
        q = re.sub(
            r"\s*Е\.\s*А\.\s*Ширяева[^А-Яа-яA-Za-z0-9]*(?:Задачник)?[^А-Яа-яA-Za-z0-9]*(?:ЕГЭбаз)?[^А-Яа-яA-Za-z0-9]*(?:2026)?[^.]*?\(тренаж[её]р\)\s*",
            " ",
            q,
            flags=re.IGNORECASE,
        )

    q = q.replace("\r", "\n")
    # Убираем инструкцию бота из конца (💡 ... )
    q = re.sub(r"\s*💡.*$", "", q, flags=re.DOTALL)

    q = re.sub(r"[∠∡]\s*([A-ZА-ЯЁ][A-ZА-ЯЁ0-9₁₂₃₄₅₆₇₈₉]*)", r"угол \1", q)
    q = q.replace("∠", "угол ").replace("∡", "угол ")

    if task.exam_type == "oge" and task.task_number == 11:
        q = re.sub(r"\s+([А-ВA-C])\)\s*y\s*=", r"\n\1) y=", q)

    if task.exam_type == "oge" and task.task_number in (20, 22) and task.image_path:
        if task.task_number == 20:
            q = re.sub(
                r"(Решите\s+(?:систему\s+уравнений|уравнение|неравенство)\s*\(см\.\s*рисунок\)).*",
                r"\1.", q, flags=re.IGNORECASE | re.DOTALL,
            )
        elif task.task_number == 22:
            q = re.sub(
                r"(Постройте\s+график\s+функции)\s+.*?\s*\(см\.\s*рисунок\)",
                r"\1 (см. рисунок)", q, flags=re.IGNORECASE | re.DOTALL,
            )
            if "(см. рисунок)" in q and "y" in q.split("(см. рисунок)")[0]:
                q = "Постройте график функции (см. рисунок)."




    # ЕГЭ база №2: соответствие величин и значений.
    # Делаем читаемый двухколоночный вид:
    # ВЕЛИЧИНЫ | ЗНАЧЕНИЯ
    # А) ...   | 1) ...
    # ...
    if task.exam_type == "ege_base" and task.task_number == 2:
        pattern = re.compile(
            r"(.*?)(?:ВЕЛИЧИНЫ)\s*"
            r"А\)\s*(.*?)\s*Б\)\s*(.*?)\s*В\)\s*(.*?)\s*Г\)\s*(.*?)\s*"
            r"(?:ЗНАЧЕНИЯ)\s*"
            r"1\)\s*(.*?)\s*2\)\s*(.*?)\s*3\)\s*(.*?)\s*4\)\s*(.*?)\s*"
            r"(В ответе.*)",
            flags=re.IGNORECASE | re.DOTALL,
        )
        m = pattern.match(q)
        if m:
            intro, a, b, c, d, v1, v2, v3, v4, tail = m.groups()
            q = (
                intro.strip()
                + "\n\nВЕЛИЧИНЫ | ЗНАЧЕНИЯ\n"
                + f"А) {a.strip()} | 1) {v1.strip()}\n"
                + f"Б) {b.strip()} | 2) {v2.strip()}\n"
                + f"В) {c.strip()} | 3) {v3.strip()}\n"
                + f"Г) {d.strip()} | 4) {v4.strip()}\n\n"
                + tail.strip()
            )
        else:
            q = re.sub(r"\s+ВЕЛИЧИНЫ\s+", "\n\nВЕЛИЧИНЫ\n", q, flags=re.IGNORECASE)
            q = re.sub(r"\s+ЗНАЧЕНИЯ\s+", "\n\nЗНАЧЕНИЯ\n", q, flags=re.IGNORECASE)
            q = re.sub(r"\s+([А-Г]\))", r"\n\1", q)
            q = re.sub(r"\s+([1-4]\))", r"\n\1", q)

    # Старый текстовый форматтер №7 отключён.
    # Для №7 используется:
    # - специальная PDF/DOCX-таблица через get_ege_base_matching_data(),
    # - либо, если вся таблица уже на картинке, остаётся только условие + изображение.


    q = q.replace("\r", "\n")

    if task.exam_type == "ege_base" and (task.task_number in (2, 7) or (task.task_number == 18 and task.image_path)):
        # Сохраняем ручные переносы строк для двухколоночного блока.
        q = re.sub(r"[ \t]+", " ", q)
        q = re.sub(r"\n\s+", "\n", q)
        q = re.sub(r"\n{3,}", "\n\n", q)
    else:
        q = re.sub(r"\s*\n\s*", " ", q)

    q = re.sub(r"\s+([,.;:!?])", r"\1", q)
    # Для заданий с сохранёнными переносами строк схлопываем только пробелы/табы,
    # но не \n — иначе А), Б), В), Г) снова сольются в одну строку.
    if task.exam_type == "ege_base" and (task.task_number in (2, 7) or (task.task_number == 18 and task.image_path)):
        q = re.sub(r"[^\S\n]{2,}", " ", q)
    else:
        q = re.sub(r"\s{2,}", " ", q)
    # Убираем переносы внутри математических выражений
    # (пробел перед - или + после цифры/скобки)
    if task.exam_type == "oge" and task.task_number in (8, 9, 13, 14, 15, 16, 17, 18, 19):
        # Схлопываем все переносы строк в пробелы
        q = re.sub(r'\s*\n\s*', ' ', q)
        q = re.sub(r'\s{2,}', ' ', q)


    return q.strip()


def answer_for_variant(task: Task) -> str:
    if task.exam_type == "oge" and task.task_number == 24:
        return "—"
    return answer_for_display(task.answer)


# ════════════════════════════════════════════════════════════════════════════
# PDF — шрифты и стили
# ════════════════════════════════════════════════════════════════════════════

def register_pdf_font() -> tuple[str, str]:
    regular_candidates = [
        "C:/Windows/Fonts/times.ttf",
        "C:/Windows/Fonts/arial.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSerif.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    ]
    bold_candidates = [
        "C:/Windows/Fonts/timesbd.ttf",
        "C:/Windows/Fonts/arialbd.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSerif-Bold.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
    ]
    regular_path = next((p for p in regular_candidates if Path(p).exists()), None)
    bold_path    = next((p for p in bold_candidates   if Path(p).exists()), None)

    if regular_path:
        pdfmetrics.registerFont(TTFont("ExamRegular", regular_path))
        regular = "ExamRegular"
    else:
        regular = "Times-Roman"

    if bold_path:
        pdfmetrics.registerFont(TTFont("ExamBold", bold_path))
        bold = "ExamBold"
    else:
        bold = "Times-Bold"

    return regular, bold


def pdf_styles() -> dict[str, ParagraphStyle]:
    regular, bold = register_pdf_font()
    base = getSampleStyleSheet()
    grey = colors.HexColor("#555555")
    return {
        "title": ParagraphStyle(
            "title", parent=base["Title"], fontName=bold, fontSize=14,
            leading=18, alignment=TA_CENTER, spaceAfter=4,
        ),
        "subtitle": ParagraphStyle(
            "subtitle", parent=base["Normal"], fontName=bold, fontSize=11,
            leading=14, alignment=TA_CENTER, spaceAfter=4,
        ),
        "small_center": ParagraphStyle(
            "small_center", parent=base["Normal"], fontName=regular, fontSize=8.5,
            leading=11, alignment=TA_CENTER, textColor=grey,
        ),
        "small_left": ParagraphStyle(
            "small_left", parent=base["Normal"], fontName=regular, fontSize=9,
            leading=12, alignment=TA_CENTER,
        ),
        "instruction": ParagraphStyle(
            "instruction", parent=base["Normal"], fontName=regular, fontSize=9.5,
            leading=13, alignment=TA_JUSTIFY,
        ),
        "part": ParagraphStyle(
            "part", parent=base["Heading2"], fontName=bold, fontSize=12,
            leading=15, alignment=TA_CENTER, spaceBefore=8, spaceAfter=5,
            borderPad=4,
        ),
        "task": ParagraphStyle(
            "task", parent=base["Normal"], fontName=regular, fontSize=10.5,
            leading=14, alignment=TA_JUSTIFY, spaceBefore=2, spaceAfter=3,
        ),
        "answer_blank": ParagraphStyle(
            "answer_blank", parent=base["Normal"], fontName=regular, fontSize=9,
            leading=12, alignment=TA_CENTER, textColor=grey,
        ),
        "footer": ParagraphStyle(
            "footer", parent=base["Normal"], fontName=regular, fontSize=8,
            leading=10, alignment=TA_CENTER, textColor=grey,
        ),
        "answer_table": ParagraphStyle(
            "answer_table", parent=base["Normal"], fontName=regular, fontSize=9,
            leading=11, alignment=TA_CENTER,
        ),
        "score_label": ParagraphStyle(
            "score_label", parent=base["Normal"], fontName=regular, fontSize=8,
            leading=10, alignment=TA_RIGHT,
            textColor=colors.HexColor("#888888"),
        ),
    }


# ════════════════════════════════════════════════════════════════════════════
# PDF — строительные блоки
# ════════════════════════════════════════════════════════════════════════════

def page_header_footer(canvas, doc, exam_title: str, variant_label: str):
    canvas.saveState()
    regular, _ = register_pdf_font()
    canvas.setFont(regular, 8)
    width, height = A4
    canvas.setFillColor(colors.HexColor("#555555"))
    canvas.drawString(17 * mm, height - 10 * mm, variant_label)
    canvas.drawRightString(width - 17 * mm, height - 10 * mm, f"{exam_title} · Кобыльская Н. Ю.")
    canvas.setStrokeColor(colors.HexColor("#cccccc"))
    canvas.line(17 * mm, height - 13 * mm, width - 17 * mm, height - 13 * mm)
    canvas.drawCentredString(width / 2, 10 * mm, f"— {doc.page} —")
    canvas.restoreState()


def add_instruction_box(story: list, styles: dict, text: str):
    box = Table(
        [[Paragraph(pdf_text(text), styles["instruction"])]],
        colWidths=[17.0 * cm],
    )
    box.setStyle(TableStyle([
        ("BOX",            (0, 0), (-1, -1), 0.7, colors.HexColor("#777777")),
        ("BACKGROUND",     (0, 0), (-1, -1), colors.HexColor("#f7f7f7")),
        ("LEFTPADDING",    (0, 0), (-1, -1), 8),
        ("RIGHTPADDING",   (0, 0), (-1, -1), 8),
        ("TOPPADDING",     (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING",  (0, 0), (-1, -1), 6),
    ]))
    story.append(box)
    story.append(Spacer(1, 4))


def make_student_header(styles: dict, variant_label: str, date_str: str) -> list:
    """
    Блок с реквизитами ученика — как на настоящем КИМ.
    """
    elems = []

    # Строка: Школа / Класс / Вариант
    top = Table(
        [[
            Paragraph("Школа: ________________________", styles["small_left"]),
            Paragraph("Класс: __________", styles["small_left"]),
            Paragraph(f"{variant_label}", styles["subtitle"]),
        ]],
        colWidths=[7.5 * cm, 4.5 * cm, 5.0 * cm],
    )
    top.setStyle(TableStyle([
        ("VALIGN",  (0, 0), (-1, -1), "BOTTOM"),
        ("ALIGN",   (2, 0), (2, 0),   "RIGHT"),
    ]))
    elems.append(top)
    elems.append(Spacer(1, 5))

    # ФИО + дата
    fio = Table(
        [[
            Paragraph("Фамилия, имя, отчество: ______________________________________________", styles["small_left"]),
            Paragraph(f"Дата: {date_str}", styles["small_left"]),
        ]],
        colWidths=[13.0 * cm, 4.0 * cm],
    )
    elems.append(fio)
    elems.append(Spacer(1, 5))

    # Горизонтальная линия-разделитель
    elems.append(HRFlowable(width="100%", thickness=0.8, color=colors.HexColor("#aaaaaa"), spaceAfter=8))
    return elems


def make_answer_boxes_pdf(n_cells: int = 8) -> Table:
    """
    Клеточки для ответа — как в настоящем КИМ.
    """
    cells = [[""] * n_cells]
    t = Table(cells, colWidths=[0.42 * cm] * n_cells, rowHeights=[0.42 * cm])
    t.setStyle(TableStyle([
        ("BOX",       (0, 0), (-1, -1), 0.8, colors.black),
        ("INNERGRID", (0, 0), (-1, -1), 0.8, colors.black),
        ("VALIGN",    (0, 0), (-1, -1), "MIDDLE"),
        ("ALIGN",     (0, 0), (-1, -1), "CENTER"),
        ("FONTSIZE",  (0, 0), (-1, -1), 11),
    ]))
    return t


def make_answer_row_pdf(styles: dict, is_part2: bool = False) -> Table:
    """
    Строка «Ответ: [клеточки]» для части 1 или пустая строка для части 2.
    """
    if is_part2:
        # Для части 2 — большие клеточки для развёрнутого решения
        rows = [[""] * 25 for _ in range(15)]
        lines = Table(
            rows,
            colWidths=[0.55 * cm] * 25,
            rowHeights=[0.55 * cm] * 15,
        )
        lines.setStyle(TableStyle([
            ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#bbbbbb")),
        ]))
        return lines

    # Часть 1 — клеточки
    label = Paragraph("<b>Ответ:</b>", styles["answer_blank"])
    boxes = make_answer_boxes_pdf(8)
    row = Table(
        [[label, boxes]],
        colWidths=[1.55 * cm, 3.9 * cm],
        hAlign="LEFT",
    )

    row.setStyle(TableStyle([
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("ALIGN", (0, 0), (-1, -1), "LEFT"),

        # Полностью убираем внутренние отступы,
        # чтобы блок ответа начинался от левого края.
        ("LEFTPADDING", (0, 0), (-1, -1), 0),
        ("RIGHTPADDING", (0, 0), (-1, -1), 0),
        ("TOPPADDING", (0, 0), (-1, -1), 0),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
    ]))

    return row




def get_ege_base_matching_data(task: Task) -> dict | None:
    """
    Достаёт данные для заданий ЕГЭ база, которые нужно выводить настоящей
    таблицей в две колонки.

    Поддерживается:
      №2 — ВЕЛИЧИНЫ / ЗНАЧЕНИЯ
      №7 — любые соответствия А–Г и 1–4:
           ПЕРИОДЫ ВРЕМЕНИ / ХАРАКТЕРИСТИКИ,
           ИНТЕРВАЛЫ / ХАРАКТЕРИСТИКИ,
           ИНТЕРВАЛЫ ВРЕМЕНИ / ХАРАКТЕРИСТИКИ,
           ГРАФИКИ / УГЛОВЫЕ КОЭФФИЦИЕНТЫ.
    """
    if task.exam_type != "ege_base" or task.task_number not in (2, 7, 18):
        return None

    q = repair_question_for_export(task) or ""
    q = q.replace("\r", "\n")
    # Убираем подсказку бота (💡 ...) до парсинга, иначе она попадёт в пункт 4).
    q = re.sub(r"\s*💡.*$", "", q, flags=re.DOTALL)
    q = re.sub(r"[ \t]+", " ", q)
    q = re.sub(r"\n+", " ", q).strip()

    # Убираем служебную строку источника.
    q = re.sub(
        r"\s*Е\.\s*А\.\s*Ширяева[^А-Яа-яA-Za-z0-9]*(?:Задачник)?[^А-Яа-яA-Za-z0-9]*(?:ЕГЭбаз)?[^А-Яа-яA-Za-z0-9]*(?:2026)?[^.]*?\(тренаж[её]р\)\s*",
        " ",
        q,
        flags=re.IGNORECASE,
    )

    # Чистим OCR-мусор в формулах вида y = kx + b и y = f(x).
    q = re.sub(r"=\s*y\s*kx\s*\+\s*b\.?", "= y = kx + b.", q, flags=re.IGNORECASE)
    q = re.sub(r"f\s*\(\s*\)\s*y\s*f\s*x\s*\.", "y = f(x).", q, flags=re.IGNORECASE)
    q = re.sub(
        r"На рисунке изображ[её]н график функции\s+y\s*=\s*f\s*\(\s*x\s*\)\s*\.",
        "На рисунке изображён график функции y = f(x).",
        q,
        flags=re.IGNORECASE,
    )
    q = re.sub(r"\s{2,}", " ", q).strip()

    def clean_tail(tail: str) -> str:
        tail = tail.strip()
        tail = re.sub(r"^Ответ:\s*А\s*Б\s*В\s*Г\s*", "", tail, flags=re.IGNORECASE)
        return tail.strip()

    # ── №2: ВЕЛИЧИНЫ / ЗНАЧЕНИЯ ──────────────────────────────────────
    if task.task_number == 2:
        pattern = re.compile(
            r"(.*?)\bВЕЛИЧИНЫ\b\s*"
            r"А\)\s*(.*?)\s*Б\)\s*(.*?)\s*В\)\s*(.*?)\s*Г\)\s*(.*?)\s*"
            r"\bЗНАЧЕНИЯ\b\s*"
            r"1\)\s*(.*?)\s*2\)\s*(.*?)\s*3\)\s*(.*?)\s*4\)\s*(.*?)\s*"
            r"(В ответе.*|Ответ:.*|$)",
            flags=re.IGNORECASE | re.DOTALL,
        )
        m = pattern.match(q)
        if not m:
            return None

        intro, a, b, c, d, v1, v2, v3, v4, tail = m.groups()
        return {
            "intro": intro.strip(),
            "left_title": "ВЕЛИЧИНЫ",
            "right_title": "ЗНАЧЕНИЯ",
            "left": [f"А) {a.strip()}", f"Б) {b.strip()}", f"В) {c.strip()}", f"Г) {d.strip()}"],
            "right": [f"1) {v1.strip()}", f"2) {v2.strip()}", f"3) {v3.strip()}", f"4) {v4.strip()}"],
            "tail": clean_tail(tail),
        }

    # ── №7: универсальная таблица соответствий ───────────────────────
    if task.task_number == 7:
        # Возвращаем None ТОЛЬКО когда ОБЕ стороны — изображения:
        # «ГРАФИКИ ФУНКЦИЙ / ГРАФИКИ ПРОИЗВОДНЫХ» и задачи с ТОЧКАМИ (латинские A B C D).
        if task.image_path and re.search(
            r"\bГРАФИКИ\s+(?:ФУНКЦИЙ|ПРОИЗВОДНЫХ)\b"
            r"|\bграфик(?:ами|ов)?\s+функц.*производн"
            r"|\bТОЧКИ\b",
            q,
            flags=re.IGNORECASE | re.DOTALL,
        ):
            return None

        # Ищем левый заголовок — без IGNORECASE: заголовки в БД хранятся
        # заглавными (ГРАФИКИ, ИНТЕРВАЛЫ), а в тексте условия слово строчное
        # («изображены графики функций»). Так избегаем ложного срабатывания.
        left_pattern = re.compile(
            r"(?<!\w)(ПЕРИОДЫ\s+ВРЕМЕНИ|ИНТЕРВАЛЫ\s+ВРЕМЕНИ|ИНТЕРВАЛЫ"
            r"|ПЕРИОДЫ|ГРАФИКИ|МЕСЯЦЫ|ФУНКЦИИ)(?!\w)",
        )
        title_match = left_pattern.search(q)
        if not title_match:
            return None

        intro = q[:title_match.start()].strip()
        left_title = re.sub(r"\s+", " ", title_match.group(1).upper()).strip()
        rest = q[title_match.end():].strip()

        # Удаляем хвост (Ответ: / В таблице...) из rest
        tail_match = re.search(r"(Ответ:.*|В таблице.*)", rest, flags=re.IGNORECASE | re.DOTALL)
        if tail_match:
            tail = rest[tail_match.start():].strip()
            rest = rest[:tail_match.start()].strip()
        else:
            tail = ""

        # Ищем правый заголовок ОТДЕЛЬНО в rest.
        # Для задач типа «ГРАФИКИ / УГЛОВЫЕ КОЭФФИЦИЕНТЫ» правый заголовок
        # стоит МЕЖДУ левыми (А, Б, В, Г) и правыми (1, 2, 3, 4) маркерами,
        # а не сразу за левым заголовком.
        right_title_re = re.compile(
            r"\b(ХАРАКТЕРИСТИКИ|УГЛОВЫЕ\s+КОЭФФИЦИЕНТЫ|КОЭФФИЦИЕНТЫ"
            r"|ЗНАЧЕНИЯ\s+ПРОИЗВОДНОЙ|ГРАФИКИ\s+ПРОИЗВОДНЫХ)\b",
            flags=re.IGNORECASE,
        )
        right_title_m = right_title_re.search(rest)
        if right_title_m:
            right_title = re.sub(r"\s+", " ", right_title_m.group(1).upper()).strip()
            left_part  = rest[:right_title_m.start()].strip()
            right_part = rest[right_title_m.end():].strip()
        else:
            right_title = "ХАРАКТЕРИСТИКИ"
            # Делим по первому числовому маркеру 1)
            first_num = re.search(r"(?<!\S)1\)", rest)
            if first_num:
                left_part  = rest[:first_num.start()].strip()
                right_part = rest[first_num.start():].strip()
            else:
                left_part  = ""
                right_part = rest

        # Парсим буквенные маркеры (А Б В Г) из left_part
        def _parse_markers(text: str, keys: str) -> dict[str, str]:
            pat = re.compile(rf"(?<!\S)([{re.escape(keys)}])\)")
            found = list(pat.finditer(text))
            result: dict[str, str] = {}
            for i, m in enumerate(found):
                key = m.group(1)
                v_start = m.end()
                v_end = found[i + 1].start() if i + 1 < len(found) else len(text)
                value = re.sub(r"\s+", " ", text[v_start:v_end]).strip()
                result[key] = value
            return result

        def _parse_num_markers(text: str) -> dict[str, str]:
            pat = re.compile(r"(?<!\S)([1-4])\)")
            found = list(pat.finditer(text))
            result: dict[str, str] = {}
            for i, m in enumerate(found):
                key = m.group(1)
                v_start = m.end()
                v_end = found[i + 1].start() if i + 1 < len(found) else len(text)
                value = re.sub(r"\s+", " ", text[v_start:v_end]).strip()
                result[key] = value
            return result

        letters = _parse_markers(left_part, "АБВГ")
        numbers = _parse_num_markers(right_part)

        # Для некоторых типов (ИНТЕРВАЛЫ) буквы содержат текст (интервал),
        # для ГРАФИКИ-типов буквы пустые (графики на картинке).
        if not all(k in letters for k in "АБВГ") or not all(k in numbers for k in "1234"):
            return None

        left = [f"{k}) {letters[k]}".rstrip() if letters.get(k) else f"{k})" for k in "АБВГ"]
        right = [f"{k}) {numbers[k]}".rstrip() if numbers.get(k) else f"{k})" for k in "1234"]

        return {
            "intro": intro,
            "left_title": left_title,
            "right_title": right_title,
            "left": left,
            "right": right,
            "tail": clean_tail(tail),
        }

    # ── №18: без изображения -> таблица соответствий ─────────────────
    if task.task_number == 18 and not task.image_path:
        work = q
        tail = ""

        # Служебные инструкции в конец таблицы не тащим.
        tail_match = re.search(
            r"(В\s+таблице.*|В\s+ответе.*|Ответ:.*)",
            work,
            flags=re.IGNORECASE | re.DOTALL,
        )
        if tail_match:
            work = work[:tail_match.start()].strip()
            tail = tail_match.group(1).strip()

        # Условие — всё до первого пункта А).
        a_match = re.search(r"(?<![А-Яа-яA-Za-z])А\)", work)
        if not a_match:
            return None

        intro = work[:a_match.start()].strip()
        intro = re.sub(r"\s+", " ", intro)
        # Убираем заголовок-пояснение вида "(А Б В Г ↔ 1 2 3 4):" из конца интро —
        # он уже понятен из заголовков таблицы.
        intro = re.sub(r"\s*\([^)]*[↔→][^)]*\)\s*:?\s*$", "", intro).strip()
        # Убираем 💡-подсказку из хвоста, если она туда просочилась.
        tail = re.sub(r"\s*💡.*$", "", tail, flags=re.DOTALL).strip()

        body = work[a_match.start():].strip()

        # Начало правого столбца — первый отдельный маркер 1).
        first_num = re.search(r"(?<!\S)1\)", body)
        if not first_num:
            return None

        left_part = body[:first_num.start()].strip()
        right_part = body[first_num.start():].strip()

        def split_items(part: str, labels: str) -> dict[str, str]:
            # Метка должна быть отдельной: А), Б), В), Г) или 1), 2), 3), 4).
            label_alt = "|".join(re.escape(ch) for ch in labels)
            item_re = re.compile(rf"(?<!\S)({label_alt})\)")
            found = list(item_re.finditer(part))
            result: dict[str, str] = {}

            for i, marker in enumerate(found):
                key = marker.group(1)
                value_start = marker.end()
                value_end = found[i + 1].start() if i + 1 < len(found) else len(part)
                value = part[value_start:value_end].strip()
                value = re.sub(r"\s+", " ", value).strip()
                result[key] = value

            return result

        letters = split_items(left_part, "АБВГ")
        numbers = split_items(right_part, "1234")

        if not all(k in letters for k in "АБВГ"):
            return None
        if not all(k in numbers for k in "1234"):
            return None

        for k in "1234":
            numbers[k] = numbers.get(k, "").strip()

        return {
            "intro": intro,
            "left_title": "НЕРАВЕНСТВА",
            "right_title": "РЕШЕНИЯ",
            "left": [
                f"А) {letters['А']}",
                f"Б) {letters['Б']}",
                f"В) {letters['В']}",
                f"Г) {letters['Г']}",
            ],
            "right": [
                f"1) {numbers['1']}",
                f"2) {numbers['2']}",
                f"3) {numbers['3']}",
                f"4) {numbers['4']}",
            ],
            "tail": tail,
        }

    return None


def make_ege_base_matching_pdf_block(task: Task, styles: dict, include_answer: bool) -> list:
    """
    Специальный блок PDF для ЕГЭ база №2 и №7:
    условие + настоящая таблица в две колонки + картинка + поле ответа.
    """
    data = get_ege_base_matching_data(task)
    if not data:
        return []

    elems = []
    score_txt = f"({task.score} балл)" if task.score == 1 else f"({task.score} балла)"

    intro = data["intro"]
    header = (
        f'<b>{task.task_number}.</b> {pdf_text(intro)} '
        f'<font size="8" color="#999999">{score_txt}</font>'
    )
    elems.append(Paragraph(header, styles["task"]))
    elems.append(Spacer(1, 4))

    table_data = [
        [
            Paragraph(f"<b>{pdf_text(data['left_title'])}</b>", styles["task"]),
            Paragraph(f"<b>{pdf_text(data['right_title'])}</b>", styles["task"]),
        ]
    ]

    for left, right in zip(data["left"], data["right"]):
        table_data.append([
            Paragraph(pdf_text(left), styles["task"]),
            Paragraph(pdf_text(right), styles["task"]),
        ])

    table = Table(table_data, colWidths=[7.6 * cm, 7.6 * cm], hAlign="LEFT")
    table.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#999999")),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#eeeeee")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("LEFTPADDING", (0, 0), (-1, -1), 5),
        ("RIGHTPADDING", (0, 0), (-1, -1), 5),
        ("TOPPADDING", (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
    ]))
    elems.append(table)
    elems.append(Spacer(1, 5))

    tail = re.sub(r"^Ответ:\s*А\s*Б\s*В\s*Г\s*", "", data["tail"], flags=re.IGNORECASE).strip()
    if tail:
        elems.append(Paragraph(pdf_text(tail), styles["task"]))
        elems.append(Spacer(1, 4))

    # Изображения
    images = resolve_image_paths(task.image_path)
    for image in images:
        try:
            img = Image(str(image))
            max_w, max_h = get_export_image_size(task, target="pdf")
            ratio = min(max_w / img.imageWidth, max_h / img.imageHeight, 1)
            img.drawWidth = img.imageWidth * ratio
            img.drawHeight = img.imageHeight * ratio
            elems.append(img)
            elems.append(Spacer(1, 3))
        except Exception:
            elems.append(Paragraph("[Изображение не удалось вставить]", styles["answer_blank"]))

    elems.append(make_answer_row_pdf(styles, is_part2=task.is_part2))
    elems.append(Spacer(1, 6))
    return elems




def make_pdf_task_block(task: Task, styles: dict, include_answer: bool) -> list:
    special = make_ege_base_matching_pdf_block(task, styles, include_answer)
    if special:
        return special

    elems = []
    q = clean_question_for_variant(task)

    # ── Убираем переносы строк для заданий с формулами ────────────────
    if (task.exam_type == "oge" and task.task_number in (8, 9, 10, 11, 13, 14, 15, 16, 17, 18, 19)) or \
       (task.exam_type == "ege_profile" and task.task_number in (6, 12)):
        q = re.sub(r'\s*\n\s*', ' ', q)
        q = re.sub(r'\s{2,}', ' ', q)

    score_txt = f"({task.score} балл)" if task.score == 1 else f"({task.score} балла)"
    q_pdf = pdf_text(q)

    if (task.exam_type == "ege_profile" and task.task_number in (6, 12)) or \
       (task.exam_type == "oge" and task.task_number in (8, 9, 13, 14, 15, 16, 17, 18, 19)):
        q_pdf = pdf_no_wrap_math_fragment(q_pdf, task.task_number)

    header = (
        f'<b>{task.task_number}.</b> {q_pdf} '
        f'<font size="8" color="#999999">{score_txt}</font>'
    )

    # ── Адаптивный размер шрифта для длинных вопросов ─────────────────
    q_len = len(q)
    if q_len > 200:
        font_size = 9.0
        leading   = 12
    elif q_len > 120:
        font_size = 9.8
        leading   = 13
    else:
        font_size = 10.5
        leading   = 14

    task_style = ParagraphStyle(
        f"task_adaptive_{task.task_number}_{q_len}",
        parent=styles["task"],
        fontSize=font_size,
        leading=leading,
        alignment=TA_JUSTIFY,
        wordWrap="LTR",
    )

    elems.append(Paragraph(header, task_style))

    # Изображения
    images = resolve_image_paths(task.image_path)
    for image in images:
        try:
            img = Image(str(image))
            max_w, max_h = get_export_image_size(task, target="pdf")
            ratio = min(max_w / img.imageWidth, max_h / img.imageHeight, 1)
            img.drawWidth  = img.imageWidth  * ratio
            img.drawHeight = img.imageHeight * ratio
            elems.append(img)
            elems.append(Spacer(1, 3))
        except Exception:
            elems.append(Paragraph("[Изображение не удалось вставить]", styles["answer_blank"]))

    # Поле ответа
    elems.append(make_answer_row_pdf(styles, is_part2=task.is_part2))
    elems.append(Spacer(1, 6))

    if task.is_part2:
        return elems
    return [KeepTogether(elems)]



def make_answers_table_pdf(tasks: list[Task], styles: dict) -> Table:
    """Таблица ответов в конце файла."""
    bold_name = "ExamBold" if "ExamBold" in pdfmetrics.getRegisteredFontNames() else "Times-Bold"

    data = [["№", "Ответ", "Балл"]]
    for task in tasks:
        ans = answer_for_variant(task)
        data.append([
            str(task.task_number),
            Paragraph(pdf_text(ans), styles["answer_table"]),
            str(task.score),
        ])

    total_score = sum(t.score for t in tasks)
    data.append(["", Paragraph("<b>Максимальный балл:</b>", styles["answer_table"]), str(total_score)])

    table = Table(data, colWidths=[1.4 * cm, 13.0 * cm, 2.5 * cm], repeatRows=1)
    table.setStyle(TableStyle([
        ("FONTNAME",       (0, 0),  (-1, 0),   bold_name),
        ("BACKGROUND",     (0, 0),  (-1, 0),   colors.HexColor("#eeeeee")),
        ("BACKGROUND",     (0, -1), (-1, -1),  colors.HexColor("#e0e0e0")),
        ("GRID",           (0, 0),  (-1, -1),  0.5, colors.HexColor("#888888")),
        ("VALIGN",         (0, 0),  (-1, -1),  "TOP"),
        ("ALIGN",          (0, 0),  (0, -1),   "CENTER"),
        ("ALIGN",          (2, 0),  (2, -1),   "CENTER"),
        ("LEFTPADDING",    (0, 0),  (-1, -1),  5),
        ("RIGHTPADDING",   (0, 0),  (-1, -1),  5),
        ("TOPPADDING",     (0, 0),  (-1, -1),  3),
        ("BOTTOMPADDING",  (0, 0),  (-1, -1),  3),
    ]))
    return table


def make_blank_part1_pdf(tasks: list[Task], styles: dict) -> list:
    """
    Бланк ответов части 1 — таблица с клеточками для каждого задания.
    Добавляется перед таблицей ответов (только для ОГЭ и ЕГЭ профиль).
    """
    part1_tasks = [t for t in tasks if not t.is_part2]
    if not part1_tasks:
        return []

    elems = []
    elems.append(Paragraph("Бланк ответов — Часть 1", styles["subtitle"]))
    elems.append(Spacer(1, 6))

    # По 4 задания в строке
    cols_per_row = 4
    rows_data = []
    row = []
    for i, task in enumerate(part1_tasks):
        cell = Table(
            [[Paragraph(f"<b>{task.task_number}</b>", styles["answer_blank"])],
             [make_answer_boxes_pdf(6)]],
            colWidths=[4.0 * cm],
        )
        cell.setStyle(TableStyle([
            ("ALIGN",  (0, 0), (-1, -1), "CENTER"),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ]))
        row.append(cell)
        if len(row) == cols_per_row:
            rows_data.append(row)
            row = []
    if row:
        # Дополняем пустыми ячейками до cols_per_row
        while len(row) < cols_per_row:
            row.append("")
        rows_data.append(row)

    blank_table = Table(rows_data, colWidths=[4.2 * cm] * cols_per_row)
    blank_table.setStyle(TableStyle([
        ("VALIGN",        (0, 0), (-1, -1), "MIDDLE"),
        ("ALIGN",         (0, 0), (-1, -1), "CENTER"),
        ("TOPPADDING",    (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ("GRID",          (0, 0), (-1, -1), 0.3, colors.HexColor("#dddddd")),
    ]))
    elems.append(blank_table)
    elems.append(Spacer(1, 10))
    return elems


# ════════════════════════════════════════════════════════════════════════════
# PDF — основная функция экспорта
# ════════════════════════════════════════════════════════════════════════════

def export_variant_pdf(
    tasks: list[Task],
    exam_type: str,
    output_path: Path,
    include_answers: bool = False,
    variant_number: int | None = None,
) -> Path:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    styles = pdf_styles()

    exam_title    = EXAM_TITLES.get(exam_type, exam_type)
    date_str      = datetime.now().strftime("%d.%m.%Y")
    year_str      = datetime.now().strftime("%Y")
    variant_label = f"Вариант № {variant_number}" if variant_number else f"Тренировочный вариант"

    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        rightMargin=17 * mm,
        leftMargin=20 * mm,
        topMargin=20 * mm,
        bottomMargin=18 * mm,
        title=f"{variant_label} · {exam_title} · {year_str}",
        author="MathBot",
    )

    story = []

    # ── Заголовок ────────────────────────────────────────────────────────────
    story.append(Paragraph(exam_title.upper(), styles["title"]))
    story.append(Paragraph(year_str, styles["small_center"]))
    story.append(Spacer(1, 8))

    # ── Реквизиты ученика ────────────────────────────────────────────────────
    story.extend(make_student_header(styles, variant_label, date_str))

    # ── Инструкция ───────────────────────────────────────────────────────────
    add_instruction_box(story, styles, get_exam_instruction(exam_type))

    # ── Справочные формулы для ЕГЭ профиль ──────────────────────────────────
    if exam_type == "ege_profile":
        story.append(Paragraph("Справочные материалы", styles["part"]))
        formulas = [
            "sin 2α = 2 sin α cos α",
            "cos 2α = cos² α − sin² α",
            "sin(α + β) = sin α cos β + cos α sin β",
            "cos(α + β) = cos α cos β − sin α sin β",
        ]
        ref_formula_style = ParagraphStyle(
            "reference_formula",
            parent=styles["instruction"],
            alignment=TA_CENTER,
            spaceBefore=0,
            spaceAfter=1,
        )
        for formula in formulas:
            story.append(Paragraph(pdf_text(formula), ref_formula_style))
        story.append(Spacer(1, 6))

    # ── Задания ──────────────────────────────────────────────────────────────
    last_part = None
    for task in tasks:
        part = get_part_title(exam_type, task.task_number)

        if part and part != last_part:
            # Часть 2 ОГЭ и ЕГЭ профиль — с новой страницы
            if last_part is not None and part == "Часть 2":
                story.append(PageBreak())

            story.append(Paragraph(part, styles["part"]))
            last_part = part

            # Инструкция к части
            if part == "Часть 1":
                if exam_type in ("oge", "ege_base"):
                    add_instruction_box(
                        story, styles,
                        "Ответом к каждому заданию является целое число или конечная десятичная дробь. "
                        "Запишите ответ в поле ответа.",
                    )
                elif exam_type == "ege_profile":
                    add_instruction_box(
                        story, styles,
                        "Ответом к заданиям 1–12 является целое число или конечная десятичная дробь. "
                        "Единицы измерения писать не нужно.",
                    )
            elif part == "Часть 2":
                if exam_type == "oge":
                    add_instruction_box(
                        story, styles,
                        "Для заданий 20–25 требуется полное обоснованное решение. "
                        "Запишите номер задания, затем решение и ответ. "
                        "Задания 20–25 оцениваются от 0 до 2 баллов каждое.",
                    )
                elif exam_type == "ege_profile":
                    add_instruction_box(
                        story, styles,
                        "Для заданий 13–19 запишите полное обоснованное решение. "
                        "Решения без обоснования не засчитываются.",
                    )

        block = make_pdf_task_block(task, styles, include_answer=include_answers)
        # make_pdf_task_block уже сам решает, нужен ли KeepTogether
        story.extend(block)

    story.append(PageBreak())

    # ── Ответы ───────────────────────────────────────────────────────────────
    if include_answers:
        story.append(Paragraph("Ответы к варианту", styles["title"]))
        story.append(Spacer(1, 8))

        # Бланк части 1 (клеточки по заданиям)
        if exam_type in ("oge", "ege_profile"):
            story.extend(make_blank_part1_pdf(tasks, styles))

        story.append(make_answers_table_pdf(tasks, styles))

    # ── Колонтитулы ──────────────────────────────────────────────────────────
    def _hf(canvas, doc_obj):
        page_header_footer(canvas, doc_obj, exam_title, variant_label)

    doc.build(story, onFirstPage=_hf, onLaterPages=_hf)

    # ── Справочные материалы в начало (ОГЭ / ЕГЭ база) ───────────────────────
    ref_pdf = find_reference_pdf(exam_type)
    if ref_pdf:
        prepend_reference_pdf(output_path, ref_pdf)

    return output_path


# ════════════════════════════════════════════════════════════════════════════
# DOCX
# ════════════════════════════════════════════════════════════════════════════

def set_docx_normal_style(doc: Document):
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)


def docx_add_run(para, text: str, bold: bool = False, italic: bool = False,
                 size_pt: int = 12, color_hex: str | None = None):
    r = para.add_run(text)
    r.bold  = bold
    r.italic = italic
    r.font.name = "Times New Roman"
    r.font.size = Pt(size_pt)
    if color_hex:
        r.font.color.rgb = RGBColor.from_string(color_hex.lstrip("#"))
    return r


def docx_add_border_para(doc: Document, text: str) -> None:
    """Абзац в рамке (инструкция)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.left_indent  = Cm(0.3)
    p.paragraph_format.right_indent = Cm(0.3)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    # Рамка через XML
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        bdr = OxmlElement(f"w:{side}")
        bdr.set(qn("w:val"), "single")
        bdr.set(qn("w:sz"), "4")
        bdr.set(qn("w:space"), "6")
        bdr.set(qn("w:color"), "888888")
        pBdr.append(bdr)
    pPr.append(pBdr)


def docx_add_shading(para, fill_hex: str = "F7F7F7"):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex)
    pPr.append(shd)


def docx_hr(doc: Document):
    """Горизонтальная линия через XML."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.space_before = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "AAAAAA")
    pBdr.append(bottom)
    pPr.append(pBdr)


def export_variant_docx(
    tasks: list[Task],
    exam_type: str,
    output_path: Path,
    include_answers: bool = False,
    variant_number: int | None = None,
) -> Path:
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    set_docx_normal_style(doc)

    section = doc.sections[0]
    section.top_margin    = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin   = Cm(2.0)
    section.right_margin  = Cm(1.5)

    exam_title    = EXAM_TITLES.get(exam_type, exam_type)
    date_str      = datetime.now().strftime("%d.%m.%Y")
    year_str      = datetime.now().strftime("%Y")
    variant_label = f"Вариант № {variant_number}" if variant_number else "Тренировочный вариант"

    # ── Заголовок ────────────────────────────────────────────────────────────
    h = doc.add_heading(f"{exam_title.upper()} · Кобыльская Н. Ю.", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h.runs[0].font.size = Pt(14)
    h.runs[0].font.name = "Times New Roman"

    yr = doc.add_paragraph(year_str)
    yr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    yr.runs[0].font.size = Pt(10)
    yr.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph()

    # ── Реквизиты ────────────────────────────────────────────────────────────
    meta_table = doc.add_table(rows=2, cols=3)
    meta_table.style = "Table Grid"
    meta_table.autofit = False

    widths = [Cm(7), Cm(4), Cm(5)]
    for i, w in enumerate(widths):
        for row in meta_table.rows:
            row.cells[i].width = w

    meta_table.cell(0, 0).text = "Школа: ____________________________"
    meta_table.cell(0, 1).text = "Класс: ____________"
    meta_table.cell(0, 2).text = variant_label
    meta_table.cell(1, 0).merge(meta_table.cell(1, 1)).text = \
        "Фамилия, имя, отчество: _________________________________________"
    meta_table.cell(1, 2).text = f"Дата: {date_str}"

    for row in meta_table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(10)
    meta_table.cell(0, 2).paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    # ── Инструкция ───────────────────────────────────────────────────────────
    docx_add_border_para(doc, get_exam_instruction(exam_type))
    doc.add_paragraph()

    # ── Справочные формулы ЕГЭ профиль ──────────────────────────────────────
    if exam_type == "ege_profile":
        ref_h = doc.add_heading("Справочные материалы", level=2)
        ref_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for formula in [
            "sin 2α = 2 sin α cos α",
            "cos 2α = cos² α − sin² α",
            "sin(α + β) = sin α cos β + cos α sin β",
            "cos(α + β) = cos α cos β − sin α sin β",
        ]:
            fp = doc.add_paragraph(formula)
            fp.paragraph_format.left_indent = Cm(1)
        doc.add_paragraph()

    # ── Задания ──────────────────────────────────────────────────────────────
    last_part = None
    for task in tasks:
        part = get_part_title(exam_type, task.task_number)

        if part and part != last_part:
            if last_part is not None and part == "Часть 2":
                doc.add_page_break()

            ph = doc.add_heading(part, level=2)
            ph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            if part == "Часть 1":
                docx_add_border_para(
                    doc,
                    "Ответом к каждому заданию является целое число или конечная десятичная дробь."
                    if exam_type in ("oge", "ege_base") else
                    "Ответом к заданиям 1–12 является целое число или конечная десятичная дробь. "
                    "Единицы измерения писать не нужно."
                )
            elif part == "Часть 2":
                docx_add_border_para(
                    doc,
                    "Для выполнения заданий этой части запишите номер задания, "
                    "затем полное обоснованное решение и ответ. "
                    "Задания оцениваются от 0 до 2 баллов каждое."
                )
            doc.add_paragraph()
            last_part = part

        # Условие задания
        score_label = f"({task.score} балл)" if task.score == 1 else f"({task.score} балла)"

        matching_data = get_ege_base_matching_data(task)
        if matching_data:
            task_para = doc.add_paragraph()
            task_para.paragraph_format.space_after = Pt(4)
            docx_add_run(task_para, f"{task.task_number}. ", bold=True, size_pt=11)
            docx_add_run(task_para, matching_data["intro"], size_pt=11)
            docx_add_run(task_para, f"  {score_label}", size_pt=8, color_hex="#999999")

            mt = doc.add_table(rows=1, cols=2)
            mt.style = "Table Grid"
            hdr = mt.rows[0].cells
            hdr[0].text = matching_data["left_title"]
            hdr[1].text = matching_data["right_title"]

            for left, right in zip(matching_data["left"], matching_data["right"]):
                row = mt.add_row().cells
                row[0].text = left
                row[1].text = right

            tail = re.sub(r"^Ответ:\s*А\s*Б\s*В\s*Г\s*", "", matching_data["tail"], flags=re.IGNORECASE).strip()
            if tail:
                doc.add_paragraph(tail)

        else:
            q = clean_question_for_variant(task)
            task_para = doc.add_paragraph()
            task_para.paragraph_format.space_after = Pt(4)

            docx_add_run(task_para, f"{task.task_number}. ", bold=True, size_pt=11)
            docx_add_run(task_para, q, size_pt=11)
            docx_add_run(task_para, f"  {score_label}", size_pt=8, color_hex="#999999")

        # Изображения
        for image in resolve_image_paths(task.image_path):
            try:
                doc.add_picture(str(image), width=get_export_image_size(task, target="docx"))
            except Exception:
                doc.add_paragraph("[Изображение не удалось вставить]")

        # Поле ответа
        if task.is_part2:
            grid = doc.add_table(rows=15, cols=25)
            grid.style = "Table Grid"
            for row in grid.rows:
                for cell in row.cells:
                    cell.width = Cm(0.5)
        else:
            ans_p = doc.add_paragraph()
            docx_add_run(ans_p, "Ответ: ", size_pt=10, color_hex="#555555")
            docx_add_run(ans_p, "□"*8, size_pt=8)

        doc.add_paragraph()

    # ── Ответы ───────────────────────────────────────────────────────────────
    if include_answers:
        doc.add_page_break()
        ah = doc.add_heading("Ответы к варианту", level=1)
        ah.alignment = WD_ALIGN_PARAGRAPH.CENTER

        ans_table = doc.add_table(rows=1, cols=3)
        ans_table.style = "Table Grid"
        hdr = ans_table.rows[0].cells
        hdr[0].text = "№"
        hdr[1].text = "Ответ"
        hdr[2].text = "Балл"
        for cell in hdr:
            for r in cell.paragraphs[0].runs:
                r.bold = True
                r.font.name = "Times New Roman"
                r.font.size = Pt(11)

        total_score = 0
        for task in tasks:
            row = ans_table.add_row().cells
            row[0].text = str(task.task_number)
            row[1].text = answer_for_variant(task)
            row[2].text = str(task.score)
            total_score += task.score
            for cell in row:
                for r in cell.paragraphs[0].runs:
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(11)

        total_row = ans_table.add_row().cells
        total_row[0].text = ""
        total_row[1].text = "Максимальный балл:"
        total_row[2].text = str(total_score)
        for cell in total_row:
            for r in cell.paragraphs[0].runs:
                r.bold = True
                r.font.name = "Times New Roman"
                r.font.size = Pt(11)

    doc.save(output_path)
    return output_path


# ════════════════════════════════════════════════════════════════════════════
# CLI и интеграция с ботом
# ════════════════════════════════════════════════════════════════════════════

def generate_variant_exports(
    exam_type: str,
    include_answers: bool = False,
    make_pdf: bool = True,
    make_docx: bool = False,
    variant_number: int | None = None,
    outdir: Path = DEFAULT_OUTDIR,
    db_path: Path = DB_PATH,
) -> list[Path]:
    tasks = build_random_variant(exam_type, db_path=db_path)
    if not tasks:
        raise RuntimeError(f"Не удалось собрать вариант: нет заданий для {exam_type}")

    outputs: list[Path] = []
    if make_pdf:
        pdf_path = outdir / make_output_name(exam_type, variant_number, include_answers, "pdf")
        outputs.append(export_variant_pdf(tasks, exam_type, pdf_path, include_answers, variant_number))

    if make_docx:
        docx_path = outdir / make_output_name(exam_type, variant_number, include_answers, "docx")
        outputs.append(export_variant_docx(tasks, exam_type, docx_path, include_answers, variant_number))

    return outputs


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Экспорт варианта ЕГЭ/ОГЭ в PDF/DOCX")
    parser.add_argument("--exam", choices=list(EXAM_RANGES.keys()), required=True)
    parser.add_argument("--pdf",    action="store_true")
    parser.add_argument("--docx",   action="store_true")
    parser.add_argument("--answers",action="store_true")
    parser.add_argument("--variant-number", type=int, default=None)
    parser.add_argument("--outdir", type=Path, default=DEFAULT_OUTDIR)
    parser.add_argument("--db",     type=Path, default=DB_PATH)
    return parser.parse_args()


# ── Функции для VK-бота (публичный API без изменений) ────────────────────────

def task_from_dict(data: dict) -> Task:
    return Task(
        id=int(data.get("id") or 0),
        exam_type=str(data.get("exam_type") or ""),
        task_number=int(data.get("task_number") or 0),
        topic=str(data.get("topic") or ""),
        question=data.get("question"),
        answer=str(data.get("answer") or ""),
        image_path=data.get("image_path"),
        set_id=data.get("set_id"),
    )


def tasks_from_dicts(items: list[dict]) -> list[Task]:
    return [task_from_dict(item) for item in items]


def export_variant_from_tasks(
    tasks: list[Task],
    exam_type: str,
    fmt: str = "pdf",
    include_answers: bool = False,
    variant_number: int | None = None,
    outdir: Path = DEFAULT_OUTDIR,
) -> Path:
    outdir.mkdir(parents=True, exist_ok=True)
    fmt = fmt.lower().strip()
    if fmt not in {"pdf", "docx"}:
        raise ValueError("fmt должен быть 'pdf' или 'docx'")
    output_path = outdir / make_output_name(exam_type, variant_number, include_answers, fmt)
    if fmt == "pdf":
        return export_variant_pdf(tasks, exam_type, output_path, include_answers, variant_number)
    return export_variant_docx(tasks, exam_type, output_path, include_answers, variant_number)


def export_variant_from_dicts(
    items: list[dict],
    exam_type: str,
    fmt: str = "pdf",
    include_answers: bool = False,
    variant_number: int | None = None,
    outdir: Path = DEFAULT_OUTDIR,
) -> Path:
    tasks = tasks_from_dicts(items)
    return export_variant_from_tasks(tasks, exam_type, fmt, include_answers, variant_number, outdir)


def main():
    args = parse_args()
    make_pdf  = args.pdf or not args.docx
    make_docx = args.docx

    outputs = generate_variant_exports(
        exam_type=args.exam,
        include_answers=args.answers,
        make_pdf=make_pdf,
        make_docx=make_docx,
        variant_number=args.variant_number,
        outdir=args.outdir,
        db_path=args.db,
    )
    print("✅ Экспорт готов:")
    for path in outputs:
        print(f"  {path}")


if __name__ == "__main__":
    main()

# NOSPLIT PATCH FOR TASKS 6 AND 12
